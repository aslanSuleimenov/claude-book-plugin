#!/usr/bin/env python3
"""
chapters_to_docx.py — Convert chapters/*.md to a formatted DOCX manuscript.

Usage:
    python chapters_to_docx.py

Requirements:
    pip install python-docx

Reads:
    - chapters/*.md (alphabetical order)
    - CLAUDE.md (for title, author, genre)

Writes:
    - versions/<title>_vNN.docx (auto-increments version)
"""

import os
import re
import glob
import sys
from pathlib import Path
from datetime import datetime


def check_dependencies():
    try:
        import docx
    except ImportError:
        print("Missing dependency: python-docx")
        print("Install with: pip install python-docx")
        sys.exit(1)


def get_project_meta() -> dict:
    """Read title and author from CLAUDE.md."""
    meta = {"title": "Untitled", "author": ""}

    claude_md = Path("CLAUDE.md")
    if not claude_md.exists():
        return meta

    content = claude_md.read_text(encoding="utf-8")

    # Extract title from first heading
    title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    if title_match:
        meta["title"] = title_match.group(1).strip()

    # Extract author if present
    author_match = re.search(r'^author:\s*(.+)$', content, re.MULTILINE | re.IGNORECASE)
    if author_match:
        meta["author"] = author_match.group(1).strip()

    return meta


def get_output_path(title: str) -> Path:
    """Generate versioned output path: versions/<title>_vNN.docx"""
    versions_dir = Path("versions")
    versions_dir.mkdir(exist_ok=True)

    # Clean title for filename
    safe_title = re.sub(r'[^\w\s-]', '', title).strip()
    safe_title = re.sub(r'[\s]+', '_', safe_title)
    if not safe_title:
        safe_title = "manuscript"

    # Find next version number
    existing = list(versions_dir.glob(f"{safe_title}_v*.docx"))
    if not existing:
        version = 1
    else:
        versions = []
        for f in existing:
            m = re.search(r'_v(\d+)\.docx$', f.name)
            if m:
                versions.append(int(m.group(1)))
        version = max(versions) + 1 if versions else 1

    return versions_dir / f"{safe_title}_v{version:02d}.docx"


def parse_markdown_chapter(content: str) -> list[dict]:
    """
    Parse markdown content into a list of blocks.
    Returns list of {"type": "heading"|"paragraph"|"break", "text": str, "level": int}
    """
    blocks = []
    lines = content.split('\n')
    current_para = []

    def flush_para():
        if current_para:
            text = ' '.join(current_para).strip()
            if text:
                blocks.append({"type": "paragraph", "text": text})
            current_para.clear()

    for line in lines:
        # Headings
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
        if heading_match:
            flush_para()
            level = len(heading_match.group(1))
            blocks.append({"type": "heading", "level": level, "text": heading_match.group(2).strip()})
            continue

        # Section break
        if re.match(r'^\s*---+\s*$', line) or re.match(r'^\s*\*\s*\*\s*\*\s*$', line):
            flush_para()
            blocks.append({"type": "break", "text": "* * *"})
            continue

        # Empty line = paragraph break
        if not line.strip():
            flush_para()
            continue

        current_para.append(line.strip())

    flush_para()
    return blocks


def apply_inline_formatting(paragraph, text: str):
    """Apply bold and italic formatting to a docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Simple pattern: **bold**, *italic*, ***bold-italic***
    pattern = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+)')
    for match in pattern.finditer(text):
        bold_italic, bold, italic, plain = match.groups()
        if bold_italic:
            run = paragraph.add_run(bold_italic)
            run.bold = True
            run.italic = True
        elif bold:
            run = paragraph.add_run(bold)
            run.bold = True
        elif italic:
            run = paragraph.add_run(italic)
            run.italic = True
        elif plain:
            paragraph.add_run(plain)


def build_docx(chapters: list[dict], meta: dict, output_path: Path):
    """Build the DOCX document."""
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION

    doc = Document()

    # Page setup: standard manuscript format
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Default paragraph style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(meta["title"])
    run.bold = True
    run.font.size = Pt(24)

    if meta["author"]:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(meta["author"])
        run.font.size = Pt(14)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(datetime.now().strftime("%B %Y"))
    run.font.size = Pt(12)

    doc.add_page_break()

    # Chapters
    first_chapter = True
    for chapter in chapters:
        if not first_chapter:
            doc.add_page_break()
        first_chapter = False

        for block in chapter["blocks"]:
            if block["type"] == "heading":
                if block["level"] == 1:
                    # Chapter heading
                    para = doc.add_heading(block["text"], level=1)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif block["level"] == 2:
                    para = doc.add_heading(block["text"], level=2)
                else:
                    para = doc.add_heading(block["text"], level=3)

            elif block["type"] == "break":
                para = doc.add_paragraph("* * *")
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para_format = para.paragraph_format
                para_format.space_before = Pt(12)
                para_format.space_after = Pt(12)

            elif block["type"] == "paragraph":
                para = doc.add_paragraph()
                para_format = para.paragraph_format
                para_format.first_line_indent = Inches(0.5)
                para_format.line_spacing = Pt(24)  # Double-spaced
                apply_inline_formatting(para, block["text"])

    doc.save(str(output_path))


def count_words(content: str) -> int:
    # Remove markdown syntax
    text = re.sub(r'#+\s', '', content)
    text = re.sub(r'\*+', '', text)
    return len(text.split())


def main():
    check_dependencies()

    # Find chapters
    chapter_files = sorted(glob.glob("chapters/*.md"))
    if not chapter_files:
        print("No chapters found in chapters/")
        print("Expected: chapters/01_title.md, chapters/02_title.md, ...")
        sys.exit(1)

    meta = get_project_meta()
    output_path = get_output_path(meta["title"])

    print(f"Title: {meta['title']}")
    print(f"Chapters: {len(chapter_files)}")
    print(f"Output: {output_path}")
    print()

    # Parse chapters
    chapters = []
    total_words = 0
    for file_path in chapter_files:
        content = Path(file_path).read_text(encoding="utf-8")
        words = count_words(content)
        total_words += words
        blocks = parse_markdown_chapter(content)
        chapters.append({"file": file_path, "blocks": blocks, "words": words})
        print(f"  {file_path} ({words:,} words)")

    print(f"\nTotal: {total_words:,} words (~{total_words // 250} min reading time)")
    print()

    # Build DOCX
    print("Building DOCX...")
    build_docx(chapters, meta, output_path)

    print(f"Done: {output_path}")


if __name__ == "__main__":
    main()
